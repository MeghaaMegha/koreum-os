"""Koreum Vault — document upload, list, delete, search, collections, lifecycle."""
import logging
import uuid
from datetime import datetime, timezone
from typing import Annotated

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status
from sqlalchemy import select

from app.core.chunker import chunk_text
from app.core.embeddings import get_embedding_provider
from app.core.vector_store import search_hybrid, search_keyword, search_vector, store_embedding
from app.database import get_db
from app.deps import CurrentUser, DBSession, require_permission
from app.models.audit import AuditEvent
from app.models.document import Document, DocumentChunk, KnowledgeCollection
from app.schemas.document import CollectionOut, DocumentOut, SearchHit, SearchResponse

logger = logging.getLogger("koreum")
router = APIRouter(prefix="/vault/documents", tags=["vault"])

MAX_UPLOAD_SIZE = 25 * 1024 * 1024  # 25 MB
ALLOWED_CONTENT_TYPES = {
    "text/plain", "text/markdown", "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/csv",
}


@router.get("", response_model=list[DocumentOut])
async def list_documents(
    current: Annotated[CurrentUser, Depends(require_permission("vault:read"))],
    db: DBSession,
    lifecycle_state: str | None = None,
    collection_id: str | None = None,
):
    query = (
        select(Document)
        .where(Document.tenant_id == current.tenant_id)
        .order_by(Document.created_at.desc())
    )
    if lifecycle_state:
        query = query.where(Document.lifecycle_state == lifecycle_state)
    if collection_id:
        query = query.where(Document.collection_id == collection_id)
    result = await db.execute(query)
    return result.scalars().all()


@router.post("", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    current: Annotated[CurrentUser, Depends(require_permission("vault:write"))],
    db: DBSession,
    file: UploadFile = File(...),
    title: str | None = None,
    collection_id: str | None = None,
):
    if file.content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file.content_type}. Allowed: {', '.join(ALLOWED_CONTENT_TYPES)}",
        )

    contents = await file.read()
    if len(contents) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 25MB)")

    # Extract text based on file type
    if file.content_type in ("text/plain", "text/markdown", "text/csv"):
        raw_text = contents.decode("utf-8", errors="ignore")
    elif file.content_type == "application/pdf":
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            raw_text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raw_text = "[PDF parsing requires PyPDF2: pip install PyPDF2]"
        except Exception:
            raw_text = "[Failed to extract PDF text]"
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import io
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(contents))
            raw_text = "\n\n".join(para.text for para in doc.paragraphs if para.text)
        except ImportError:
            raw_text = "[DOCX parsing requires python-docx: pip install python-docx]"
        except Exception:
            raw_text = "[Failed to extract DOCX text]"
    else:
        raw_text = "[Unsupported file type for text extraction]"

    # Validate collection if provided
    collection_uuid = None
    if collection_id:
        col_result = await db.execute(
            select(KnowledgeCollection).where(
                KnowledgeCollection.id == collection_id,
                KnowledgeCollection.tenant_id == current.tenant_id,
            )
        )
        if not col_result.scalar_one_or_none():
            raise HTTPException(status_code=404, detail="Collection not found")
        collection_uuid = uuid.UUID(collection_id)

    # Create document
    doc = Document(
        tenant_id=current.tenant_id,
        uploaded_by=current.id,
        title=title or file.filename or "Untitled",
        filename=file.filename or "unknown",
        content_type=file.content_type or "application/octet-stream",
        file_size=len(contents),
        status="processing",
        raw_text=raw_text,
        collection_id=collection_uuid,
        version=1,
        lifecycle_state="active",
        source_type="upload",
    )
    db.add(doc)
    await db.flush()

    # Chunk the text
    chunks = chunk_text(raw_text) if raw_text else []

    # Create chunk records
    chunk_records = []
    for chunk in chunks:
        chunk_record = DocumentChunk(
            document_id=doc.id,
            chunk_index=chunk.index,
            content=chunk.content,
        )
        db.add(chunk_record)
        chunk_records.append(chunk_record)
    await db.flush()

    # Generate and store embeddings
    if chunk_records:
        try:
            provider = get_embedding_provider()
            texts = [c.content for c in chunk_records]
            embeddings = provider.embed_batch(texts)
            for chunk_record, embedding in zip(chunk_records, embeddings):
                await store_embedding(db, chunk_record.id, embedding)
            doc.status = "indexed"
            logger.info(f"Indexed document {doc.id} with {len(chunks)} chunks")
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            doc.status = "uploaded"
    else:
        doc.status = "uploaded"

    db.add(
        AuditEvent(
            tenant_id=current.tenant_id,
            actor_user_id=current.id,
            action="VAULT_DOCUMENT_UPLOAD",
            details={"document_id": str(doc.id), "filename": doc.filename, "chunks": len(chunks), "version": 1},
            created_at=datetime.now(timezone.utc),
        )
    )
    await db.commit()
    await db.refresh(doc)
    return doc


@router.get("/{document_id}", response_model=DocumentOut)
async def get_document(
    document_id: str,
    current: Annotated[CurrentUser, Depends(require_permission("vault:read"))],
    db: DBSession,
):
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.tenant_id == current.tenant_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.put("/{document_id}/lifecycle", response_model=DocumentOut)
async def update_lifecycle(
    document_id: str,
    current: Annotated[CurrentUser, Depends(require_permission("vault:write"))],
    db: DBSession,
    state: str = Query("active", regex="^(draft|active|archived)$"),
):
    """Update document lifecycle state: draft, active, or archived."""
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.tenant_id == current.tenant_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    old_state = doc.lifecycle_state
    doc.lifecycle_state = state

    db.add(
        AuditEvent(
            tenant_id=current.tenant_id,
            actor_user_id=current.id,
            action="VAULT_DOCUMENT_LIFECYCLE",
            details={"document_id": str(doc.id), "old_state": old_state, "new_state": state},
            created_at=datetime.now(timezone.utc),
        )
    )
    await db.commit()
    await db.refresh(doc)
    return doc


@router.post("/{document_id}/version", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def create_version(
    document_id: str,
    current: Annotated[CurrentUser, Depends(require_permission("vault:write"))],
    db: DBSession,
    file: UploadFile = File(...),
    title: str | None = None,
):
    """Create a new version of an existing document."""
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.tenant_id == current.tenant_id)
    )
    parent_doc = result.scalar_one_or_none()
    if not parent_doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Archive old version
    parent_doc.lifecycle_state = "archived"

    # Read new file
    contents = await file.read()
    if file.content_type in ("text/plain", "text/markdown", "text/csv"):
        raw_text = contents.decode("utf-8", errors="ignore")
    elif file.content_type == "application/pdf":
        try:
            import io
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(contents))
            raw_text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            raw_text = "[Failed to extract PDF text]"
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import io
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(io.BytesIO(contents))
            raw_text = "\n\n".join(para.text for para in docx_doc.paragraphs if para.text)
        except Exception:
            raw_text = "[Failed to extract DOCX text]"
    else:
        raw_text = "[Unsupported file type]"

    # Create new version
    new_doc = Document(
        tenant_id=current.tenant_id,
        uploaded_by=current.id,
        title=title or parent_doc.title,
        filename=file.filename or parent_doc.filename,
        content_type=file.content_type or "application/octet-stream",
        file_size=len(contents),
        status="processing",
        raw_text=raw_text,
        collection_id=parent_doc.collection_id,
        version=parent_doc.version + 1,
        parent_document_id=parent_doc.id,
        lifecycle_state="active",
        source_type="upload",
    )
    db.add(new_doc)
    await db.flush()

    # Chunk and embed
    chunks = chunk_text(raw_text) if raw_text else []
    chunk_records = []
    for chunk in chunks:
        chunk_record = DocumentChunk(
            document_id=new_doc.id,
            chunk_index=chunk.index,
            content=chunk.content,
        )
        db.add(chunk_record)
        chunk_records.append(chunk_record)
    await db.flush()

    if chunk_records:
        try:
            provider = get_embedding_provider()
            texts = [c.content for c in chunk_records]
            embeddings = provider.embed_batch(texts)
            for chunk_record, embedding in zip(chunk_records, embeddings):
                await store_embedding(db, chunk_record.id, embedding)
            new_doc.status = "indexed"
            logger.info(f"Indexed document {new_doc.id} v{new_doc.version} with {len(chunks)} chunks")
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            new_doc.status = "uploaded"

    db.add(
        AuditEvent(
            tenant_id=current.tenant_id,
            actor_user_id=current.id,
            action="VAULT_DOCUMENT_VERSION",
            details={"document_id": str(new_doc.id), "parent_id": str(parent_doc.id), "version": new_doc.version},
            created_at=datetime.now(timezone.utc),
        )
    )
    await db.commit()
    await db.refresh(new_doc)
    return new_doc


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: str,
    current: Annotated[CurrentUser, Depends(require_permission("vault:delete"))],
    db: DBSession,
):
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.tenant_id == current.tenant_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.delete(doc)

    db.add(
        AuditEvent(
            tenant_id=current.tenant_id,
            actor_user_id=current.id,
            action="VAULT_DOCUMENT_DELETE",
            details={"document_id": str(document_id)},
            created_at=datetime.now(timezone.utc),
        )
    )
    await db.commit()
    return None


@router.post("/search", response_model=SearchResponse)
async def search_documents(
    current: Annotated[CurrentUser, Depends(require_permission("vault:read"))],
    db: DBSession,
    query: str = "",
    limit: int = 10,
    mode: str = Query("hybrid", regex="^(vector|keyword|hybrid)$"),
):
    """Semantic search through document chunks.

    Modes:
    - vector: pure vector similarity search
    - keyword: pure keyword (ILIKE) search
    - hybrid: combined vector + keyword search (default)
    """
    if not query.strip():
        return SearchResponse(query=query, total=0, hits=[], confidence=0.0, evidence=[])

    try:
        if mode == "keyword":
            hits = await search_keyword(db, current.tenant_id, query, limit)
        elif mode == "vector":
            provider = get_embedding_provider()
            query_embedding = provider.embed(query)
            hits = await search_vector(db, current.tenant_id, query_embedding, limit)
        else:  # hybrid
            provider = get_embedding_provider()
            query_embedding = provider.embed(query)
            hits = await search_hybrid(db, current.tenant_id, query, query_embedding, limit)

        # Compute confidence based on top hit score
        confidence = 0.0
        if hits:
            top_score = hits[0]["score"]
            confidence = min(top_score, 1.0)

        # Build evidence summary
        evidence = []
        seen_docs = set()
        for hit in hits:
            doc_id = hit["document_id"]
            if doc_id not in seen_docs:
                evidence.append({
                    "document_id": doc_id,
                    "document_title": hit["document_title"],
                    "source_citation": hit["source_citation"],
                    "score": round(hit["score"], 4),
                    "chunk_count": sum(1 for h in hits if h["document_id"] == doc_id),
                })
                seen_docs.add(doc_id)

        return SearchResponse(
            query=query,
            total=len(hits),
            hits=[SearchHit(**hit) for hit in hits],
            confidence=round(confidence, 4),
            evidence=evidence,
        )
    except Exception as e:
        logger.error(f"Search failed ({mode}): {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")


# ─── Knowledge Collections ───

@router.get("/collections/all", response_model=list[CollectionOut])
async def list_collections(
    current: Annotated[CurrentUser, Depends(require_permission("vault:read"))],
    db: DBSession,
):
    result = await db.execute(
        select(KnowledgeCollection)
        .where(KnowledgeCollection.tenant_id == current.tenant_id)
        .order_by(KnowledgeCollection.created_at.desc())
    )
    collections = result.scalars().all()

    # Count documents in each collection
    out = []
    for col in collections:
        doc_count_result = await db.execute(
            select(Document).where(Document.collection_id == col.id, Document.lifecycle_state == "active")
        )
        out.append(
            CollectionOut(
                id=col.id,
                tenant_id=col.tenant_id,
                name=col.name,
                description=col.description,
                created_at=col.created_at,
                document_count=len(doc_count_result.scalars().all()),
            )
        )
    return out


@router.post("/collections", response_model=CollectionOut, status_code=status.HTTP_201_CREATED)
async def create_collection(
    current: Annotated[CurrentUser, Depends(require_permission("vault:write"))],
    db: DBSession,
    name: str = Query(..., min_length=1),
    description: str | None = None,
):
    col = KnowledgeCollection(
        tenant_id=current.tenant_id,
        name=name,
        description=description,
    )
    db.add(col)

    db.add(
        AuditEvent(
            tenant_id=current.tenant_id,
            actor_user_id=current.id,
            action="VAULT_COLLECTION_CREATE",
            details={"collection_id": str(col.id), "name": name},
            created_at=datetime.now(timezone.utc),
        )
    )
    await db.commit()
    await db.refresh(col)
    return CollectionOut(
        id=col.id,
        tenant_id=col.tenant_id,
        name=col.name,
        description=col.description,
        created_at=col.created_at,
        document_count=0,
    )
